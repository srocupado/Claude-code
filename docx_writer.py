import os
import logging
from datetime import date, timedelta

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

OUTPUT_DIR   = "output"
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template.docx")

COLOR_TEXT  = RGBColor(0x33, 0x33, 0x33)   # #333333 — template color
COLOR_RED   = RGBColor(0xFF, 0x00, 0x00)   # emendas deadline
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)   # table label cells


def _set_margins(doc: Document, top=2.0, bottom=2.0, left=3.0, right=1.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def _set_default_font(doc: Document):
    """Template base style: Times New Roman 10 pt, no explicit color."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)


# ── Low-level helpers ────────────────────────────────────────────────────────

def _new_para(doc: Document, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Create a paragraph with zero space_before / space_after."""
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    return para


def _styled_run(para, text: str, *, bold=False, size=12, font_name: str | None = None):
    """Add a run with explicit color and optional bold/font/size."""
    r = para.add_run(text)
    r.bold           = bold
    r.font.size      = Pt(size)
    r.font.color.rgb = COLOR_TEXT
    if font_name:
        r.font.name = font_name
    return r


def _blank(doc: Document):
    """Empty paragraph — vertical spacer with zero spacing."""
    _new_para(doc)


def _add_divider(doc: Document):
    """Paragraph with bottom border in template color (#333333)."""
    para = _new_para(doc)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "333333")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Prazos table ─────────────────────────────────────────────────────────────

def _add_prazos_table(doc: Document, pub_date: date) -> date:
    """3×2 deadline table positioned to the right side of the page.

      Eficácia:      DD/MM/YYYY a DD/MM/YYYY, prorrogável por mais 60 dias
      Sobrestamento: DD/MM/YYYY
      Emendas:       DD/MM/YYYY a DD/MM/YYYY   ← bold + red
    """
    eficacia_end  = pub_date + timedelta(days=59)   # day 60 (day 1 = publication)
    sobrestamento = pub_date + timedelta(days=45)
    emendas_end   = pub_date + timedelta(days=7)

    table = doc.add_table(rows=3, cols=2)
    tbl   = table._tbl

    # ── Table-level XML properties ────────────────────────────────────────────
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5081")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "3856")   # ≈ 6.8 cm indent → table on right side
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    tblBorders = OxmlElement("w:tblBorders")
    for bname in ("insideH", "insideV"):
        b = OxmlElement(f"w:{bname}")
        b.set(qn("w:val"),   "dotted")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # ── Column grid (widths) ──────────────────────────────────────────────────
    tblGrid = OxmlElement("w:tblGrid")
    for w in (1956, 3125):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tblPr_idx = list(tbl).index(tblPr)
    tbl.insert(tblPr_idx + 1, tblGrid)

    # ── Row data ──────────────────────────────────────────────────────────────
    rows_data = [
        (
            "Eficácia:",
            f"{pub_date.strftime('%d/%m/%Y')} a {eficacia_end.strftime('%d/%m/%Y')},"
            f" prorrogável por mais 60 dias",
            False, COLOR_BLACK,
        ),
        (
            "Sobrestamento:",
            sobrestamento.strftime("%d/%m/%Y"),
            False, COLOR_BLACK,
        ),
        (
            "Emendas:",
            f"{pub_date.strftime('%d/%m/%Y')} a {emendas_end.strftime('%d/%m/%Y')}",
            True, COLOR_RED,   # bold + red — action deadline
        ),
    ]

    for row_idx, (label, value, bold_val, val_color) in enumerate(rows_data):
        row = table.rows[row_idx]
        tr  = row._tr

        # Exact row height: 571 dxa
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            tr.insert(0, trPr)
        trH = OxmlElement("w:trHeight")
        trH.set(qn("w:hRule"), "exact")
        trH.set(qn("w:val"),   "571")
        trPr.append(trH)

        # Cell widths
        for cell, width in zip(row.cells, (1956, 3125)):
            tc   = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"),    str(width))
            tcW.set(qn("w:type"), "dxa")
            tcPr.insert(0, tcW)

        # Label cell (col 0)
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p0.paragraph_format.space_after = Pt(0)
        rl = p0.add_run(label)
        rl.font.name      = "Arial"
        rl.font.color.rgb = COLOR_BLACK

        # Value cell (col 1)
        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p1.paragraph_format.space_after = Pt(0)
        rv = p1.add_run(value)
        rv.font.name      = "Arial"
        rv.bold           = bold_val
        rv.font.color.rgb = val_color

    return emendas_end


def _add_atencao(doc: Document, emendas_end: date):
    """Red bold notice about amendment submission deadline."""
    para = _new_para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    text = (
        f"A T E N Ç Ã O !  "
        f"AS EMENDAS DEVERÃO SER ENVIADAS PELO INFOLEG-AUTENTICADOR "
        f"ATÉ 23h59min DO DIA {emendas_end.strftime('%d/%m/%Y')}."
    )
    r = para.add_run(text)
    r.bold           = True
    r.font.size      = Pt(10)
    r.font.color.rgb = COLOR_RED
    r.font.name      = "Arial"


# ── Document-level builders ──────────────────────────────────────────────────

def _add_title(doc: Document, text: str):
    para = _new_para(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _styled_run(para, text, bold=True, size=14, font_name="Source Sans Pro")


def _add_subtitle(doc: Document, text: str):
    para = _new_para(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _styled_run(para, text, bold=True, size=14, font_name="Source Sans Pro")


def _add_metadata_line(doc: Document, label: str, value: str):
    para = _new_para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _styled_run(para, label + " ", bold=True,  size=12)
    _styled_run(para, value,        bold=False, size=12)


def _add_section_heading(doc: Document, text: str):
    """Blank line → bold heading → blank line."""
    _blank(doc)
    para = _new_para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _styled_run(para, text, bold=True, size=12)
    _blank(doc)


def _add_labeled_block(doc: Document, label: str, body: str):
    """Blank → bold label → blank → body paragraphs."""
    _blank(doc)
    para = _new_para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _styled_run(para, label, bold=True, size=12)
    _blank(doc)
    _add_body_text(doc, body)


def _add_body_text(doc: Document, text: str):
    """Split on double newlines → one paragraph; single newline → line break."""
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [text.strip()] if text.strip() else []
    for para_text in paragraphs:
        lines = para_text.splitlines()
        para = _new_para(doc, WD_ALIGN_PARAGRAPH.JUSTIFY)
        for j, line in enumerate(lines):
            r = para.add_run(line)
            r.font.size      = Pt(12)
            r.font.color.rgb = COLOR_TEXT
            if j < len(lines) - 1:
                r.add_break()


# ── Public API ───────────────────────────────────────────────────────────────

def write_nota_tecnica(mp: dict, content: dict, output_dir: str = OUTPUT_DIR) -> str:
    os.makedirs(output_dir, exist_ok=True)

    pub_date  = date.fromisoformat(mp["data_publicacao"]) if mp.get("data_publicacao") else date.today()
    prazo_60  = pub_date + timedelta(days=60)
    prazo_120 = pub_date + timedelta(days=120)

    # Open template as base — inherits all styles, fonts and theme
    doc = Document(TEMPLATE_PATH) if os.path.exists(TEMPLATE_PATH) else Document()
    _set_margins(doc)
    _set_default_font(doc)

    # ── Prazos table (first element — before title) ───────────────────────────
    # python-docx Document() has no default paragraph, just sectPr.
    # add_table() inserts before sectPr → table ends up at body[0]. ✓
    emendas_end = _add_prazos_table(doc, pub_date)
    _add_atencao(doc, emendas_end)
    _blank(doc)

    # ── Title & subtitle ──────────────────────────────────────────────────────
    title    = content.get("titulo",    f"NOTA TÉCNICA MP nº {mp['numero']}/{mp['ano']}")
    subtitle = content.get("subtitulo", "Análise de Impacto da Medida Provisória")
    _add_title(doc, title)
    _add_subtitle(doc, subtitle)
    _add_divider(doc)

    # ── Metadata ──────────────────────────────────────────────────────────────
    _add_metadata_line(doc, "Expedidor:", "Poder Executivo – Presidência da República")
    _add_metadata_line(doc, "Publicação no DOU (Edição Extra):", pub_date.strftime("%d/%m/%Y"))
    _add_metadata_line(doc, "Vigência imediata (art. 62, §3º, CF):", pub_date.strftime("%d/%m/%Y"))
    _add_metadata_line(doc, "Prazo de vigência – 1ª prorrogação (60 dias):", prazo_60.strftime("%d/%m/%Y"))
    _add_metadata_line(doc, "Prazo máximo de vigência – 2ª prorrogação (120 dias):", prazo_120.strftime("%d/%m/%Y"))
    _add_metadata_line(doc, "Tramitação:", "Comissão Mista → Câmara dos Deputados → Senado Federal")
    _add_metadata_line(doc, "Relator na comissão mista:", "a designar")
    _add_metadata_line(doc, "Data de atualização:", date.today().strftime("%d/%m/%Y"))
    if mp.get("url_planalto"):
        _add_metadata_line(doc, "Texto no Planalto:", mp["url_planalto"])

    # ── Ementa / Explicação da matéria ────────────────────────────────────────
    _add_labeled_block(
        doc,
        "Ementa / Explicação da matéria:",
        content.get("ementa_expandida", mp.get("ementa", "")),
    )

    # ── Numbered sections ─────────────────────────────────────────────────────
    for i in range(1, 7):
        title_key = f"secao_{i}_titulo"
        body_key  = f"secao_{i}_conteudo"
        if title_key in content:
            _add_section_heading(doc, content[title_key])
            if body_key in content:
                _add_body_text(doc, content[body_key])

    # ── Arguments and recommendation ─────────────────────────────────────────
    for label, key in [
        ("Argumento favorável:",      "argumento_favoravel"),
        ("Argumento contrário:",      "argumento_contrario"),
        ("Recomendação estratégica:", "recomendacao"),
    ]:
        if content.get(key):
            _add_labeled_block(doc, label, content[key])

    # ── Fixed closing block (hardcoded — not generated by AI) ─────────────────
    _blank(doc)
    _add_divider(doc)
    _blank(doc)
    _add_metadata_line(
        doc,
        "Vigência:",
        "A Medida Provisória entra em vigor na data de sua publicação.",
    )
    _blank(doc)
    sig = _new_para(doc, WD_ALIGN_PARAGRAPH.LEFT)
    _styled_run(sig, "Assessoria da Liderança do Podemos", bold=True, size=12)

    # ── Save ──────────────────────────────────────────────────────────────────
    # ASCII-only filename: special chars (É, º) break GitHub Actions artifact ZIP
    filename = f"NOTA_TECNICA_-_MPV_n{mp['numero']}_de_{mp['ano']}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    logger.info("Nota técnica salva: %s", filepath)
    return filepath
